"""Build the Structural Fuzzing book as a professional .docx.

Usage:
    python build_book.py [--skip-illustrations] [--skip-pandoc] [--skip-postprocess]
"""

import argparse
import re
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor

BOOK_DIR = Path(__file__).parent
IMAGES_DIR = BOOK_DIR / "images"
OUTPUT_DIR = BOOK_DIR / "output"

# Chapter ordering with part markers
# Tuples = (filename, part_title) for first chapter in each part
CHAPTERS = [
    # Part I: Foundations
    ("part1-foundations/chapter-01-why-geometry.md", "Part I: Foundations"),
    "part1-foundations/chapter-02-mahalanobis-distance.md",
    "part1-foundations/chapter-03-hyperbolic-geometry.md",
    "part1-foundations/chapter-04-spd-manifolds.md",
    "part1-foundations/chapter-05-topological-data-analysis.md",
    "part1-foundations/chapter-5b-spectral-geometry-and-the-angular-basis.md",
    # Part II: Algorithms
    ("part2-algorithms/chapter-06-pathfinding-on-manifolds.md", "Part II: Algorithms"),
    "part2-algorithms/chapter-07-equilibrium-on-manifolds.md",
    "part2-algorithms/chapter-08-pareto-optimization.md",
    "part2-algorithms/chapter-09-adversarial-robustness.md",
    "part2-algorithms/chapter-10-adversarial-probing.md",
    # Part III: Design Patterns
    ("part3-patterns/chapter-11-subset-enumeration.md", "Part III: Design Patterns"),
    "part3-patterns/chapter-12-compositional-testing.md",
    "part3-patterns/chapter-13-group-theoretic-augmentation.md",
    "part3-patterns/chapter-14-gradient-reversal.md",
    "part3-patterns/chapter-15-cholesky-parameterization.md",
    # Part IV: Systems
    ("part4-systems/chapter-16-geometric-pipelines.md", "Part IV: Systems"),
    "part4-systems/chapter-17-scaling.md",
    "part4-systems/chapter-18-production-deployment.md",
    "part4-systems/chapter-19-case-study-defect-prediction.md",
    "part4-systems/chapter-20-case-study-bioacoustics.md",
    "part4-systems/chapter-21-case-study-aesthetic-judgment.md",
    "part4-systems/chapter-22-case-study-legal-embeddings.md",
    # Appendices
    ("appendices/appendix-a-notation.md", "Appendices"),
    "appendices/appendix-b-software.md",
    "appendices/appendix-c-proofs.md",
]

# Map chapter number to illustration filename
ILLUSTRATIONS = {
    1: (
        "ch01-geometric-toolchain.png",
        "The Structural Fuzzing pipeline: six stages built on geometric foundations.",
    ),
    2: (
        "ch02-euclidean-vs-mahalanobis.png",
        "Isodistance contours: Euclidean circles vs Mahalanobis ellipses.",
    ),
    3: (
        "ch03-poincare-ball.png",
        "A hierarchical tree embedded in the Poincar\u00e9 ball, with depth increasing toward the boundary.",
    ),
    4: (
        "ch04-spd-manifold.png",
        "SPD matrices visualized as ellipses (left) unfold into flat log-Euclidean space (right).",
    ),
    5: (
        "ch05-tda-persistence.png",
        "Vietoris\u2013Rips filtration at increasing radii, with the resulting persistence diagram.",
    ),
    6: (
        "ch06-pathfinding.png",
        "A* pathfinding on a decision manifold: the manifold-aware path respects the moral boundary.",
    ),
    7: (
        "ch07-nash-vs-bge.png",
        "Ultimatum game: Nash equilibrium maximizes scalar payoff; the Bond Geodesic Equilibrium balances multiple dimensions.",
    ),
    8: (
        "ch08-pareto-frontier.png",
        "Pareto frontier in dimension-count vs prediction-error space. Gold points are non-dominated.",
    ),
    9: (
        "ch09-robustness.png",
        "Left: broad vs sharp loss-landscape minima. Right: MRI perturbation distribution with percentile markers.",
    ),
    11: (
        "ch11-hasse-diagram.png",
        "Hasse diagram of the subset lattice for four feature dimensions, colored by prediction error.",
    ),
    12: (
        "ch12-interaction-heatmap.png",
        "Pairwise dimension interaction matrix: red indicates synergy, blue indicates redundancy.",
    ),
    13: (
        "ch13-dihedral-grid.png",
        "All eight symmetries of the dihedral group D\u2084 applied to an L-shaped pattern.",
    ),
    16: (
        "ch16-pipeline-architecture.png",
        "Six-stage structural fuzzing pipeline with data flow to the domain-specific evaluation function.",
    ),
    20: (
        "ch20-bioacoustics.png",
        "Sperm whale coda spectrogram (left) and cetacean taxonomy on the Poincar\u00e9 disk (right).",
    ),
}


# --- chapter numbers are references, resolved by manifest position -------------
_ORIG_NUM_TO_KEY = {
    1: "why-geometry",
    2: "mahalanobis-distance",
    3: "hyperbolic-geometry",
    4: "spd-manifolds",
    5: "topological-data-analysis",
    6: "pathfinding-on-manifolds",
    7: "equilibrium-on-manifolds",
    8: "pareto-optimization",
    9: "adversarial-robustness",
    11: "subset-enumeration",
    12: "compositional-testing",
    13: "group-theoretic-augmentation",
    16: "geometric-pipelines",
    20: "case-study-bioacoustics",
}
# illustrations keyed by chapter slug (decoupled from the now-positional number)
ILLUSTRATIONS_BY_KEY = {_ORIG_NUM_TO_KEY[n]: v for n, v in ILLUSTRATIONS.items()}

_TOKEN = re.compile(r"\{\{ch:([a-z0-9-]+)\}\}")


def _chapter_key(filename):
    """Filename slug -> stable chapter key (chapter-03-hyperbolic-geometry -> that)."""
    m = re.match(r"chapter-\d+[a-z]?-(.+)\.md$", Path(filename).name)
    return m.group(1) if m else None


def _number_by_key():
    """Assign chapter numbers by manifest position (appendices excluded)."""
    num, i = {}, 0
    for entry in CHAPTERS:
        fn = entry[0] if isinstance(entry, tuple) else entry
        k = _chapter_key(fn)
        if k:
            i += 1
            num[k] = i
    return num


NUMBER_BY_KEY = _number_by_key()


def _resolve_tokens(text):
    """Replace {{ch:key}} with the chapter number from its manifest position."""
    return _TOKEN.sub(
        lambda m: str(NUMBER_BY_KEY.get(m.group(1), f"?{m.group(1)}?")), text
    )


def _extract_chapter_num(filename):
    """Extract chapter number from filename like 'chapter-06-...' or 'appendix-a-...'."""
    stem = Path(filename).stem
    parts = stem.split("-")
    if parts[0] == "chapter":
        try:
            return int(parts[1])
        except ValueError:
            return None
    return None


def create_reference_docx():
    """Generate reference.docx programmatically for pandoc styling."""
    ref_path = BOOK_DIR / "reference.docx"
    doc = Document()

    # Page setup
    for section in doc.sections:
        section.page_height = Inches(11)
        section.page_width = Inches(8.5)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # Normal style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    # Heading styles
    for level, size, color, space_before in [
        ("Heading 1", 24, RGBColor(0x1A, 0x1A, 0x2E), 36),
        ("Heading 2", 16, RGBColor(0x2C, 0x3E, 0x6B), 18),
        ("Heading 3", 13, RGBColor(0x2C, 0x3E, 0x6B), 12),
        ("Heading 4", 11, RGBColor(0x2C, 0x3E, 0x6B), 10),
    ]:
        h = doc.styles[level]
        h.font.name = "Calibri"
        h.font.size = Pt(size)
        h.font.color.rgb = color
        h.font.bold = True
        h.paragraph_format.space_before = Pt(space_before)
        h.paragraph_format.space_after = Pt(8)
        if level == "Heading 1":
            h.paragraph_format.page_break_before = True

    # Add placeholder paragraphs in each style so pandoc picks them up
    doc.add_heading("Heading 1", level=1)
    doc.add_heading("Heading 2", level=2)
    doc.add_heading("Heading 3", level=3)
    p = doc.add_paragraph("Normal text")
    # Source Code style for code blocks
    try:
        sc = doc.styles.add_style("Source Code", 1)  # 1 = paragraph style
        sc.font.name = "Consolas"
        sc.font.size = Pt(9)
        sc.paragraph_format.space_before = Pt(2)
        sc.paragraph_format.space_after = Pt(2)
        sc.paragraph_format.line_spacing = 1.0
    except Exception:
        pass  # Style may already exist

    # Block Text for blockquotes
    try:
        bt = doc.styles["Block Text"]
        bt.font.size = Pt(10.5)
        bt.font.italic = True
        bt.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        bt.paragraph_format.left_indent = Inches(0.5)
        bt.paragraph_format.right_indent = Inches(0.5)
    except Exception:
        pass

    doc.save(str(ref_path))
    print(f"Created reference.docx: {ref_path}")
    return ref_path


def concatenate_chapters():
    """Concatenate all chapter markdown with metadata, parts, and image refs."""
    combined_path = BOOK_DIR / "output" / "_combined.md"

    lines = []

    # YAML metadata
    lines.append("---")
    lines.append('title: "Structural Fuzzing"')
    lines.append('subtitle: "Geometric Methods for Adversarial Model Validation"')
    lines.append('author: "Andrew H. Bond"')
    lines.append('date: "2026"')
    lines.append("---")
    lines.append("")
    lines.append("\\newpage")
    lines.append("")

    for entry in CHAPTERS:
        if isinstance(entry, tuple):
            filename, part_title = entry
        else:
            filename = entry
            part_title = None

        # Insert part divider
        if part_title:
            lines.append("\\newpage")
            lines.append("")
            lines.append(f"# {part_title}")
            lines.append("")
            lines.append("\\newpage")
            lines.append("")

        # Read chapter content
        chapter_path = BOOK_DIR / filename
        if not chapter_path.exists():
            print(f"WARNING: {chapter_path} not found, skipping")
            continue

        content = chapter_path.read_text(encoding="utf-8")
        content = _resolve_tokens(content)  # {{ch:key}} -> manifest-position number

        # Insert illustration after first section heading (## x.x)
        ch_key = _chapter_key(filename)
        if ch_key and ch_key in ILLUSTRATIONS_BY_KEY:
            img_file, caption = ILLUSTRATIONS_BY_KEY[ch_key]
            img_path = IMAGES_DIR / img_file
            if img_path.exists():
                # Find the first ## heading and insert illustration before it
                content_lines = content.split("\n")
                inserted = False
                new_lines = []
                for line in content_lines:
                    if not inserted and line.startswith("## "):
                        # Insert figure before first section
                        new_lines.append("")
                        new_lines.append(f"![{caption}](images/{img_file})")
                        new_lines.append("")
                        inserted = True
                    new_lines.append(line)
                if not inserted:
                    # Fallback: insert at end of chapter
                    new_lines.append("")
                    new_lines.append(f"![{caption}](images/{img_file})")
                    new_lines.append("")
                content = "\n".join(new_lines)

        lines.append(content)
        lines.append("")
        lines.append("\\newpage")
        lines.append("")

    combined_path.write_text("\n".join(lines), encoding="utf-8")
    print(f"Combined markdown: {combined_path} ({len(lines)} lines)")
    return combined_path


def run_pandoc(combined_md):
    """Convert combined markdown to docx via pandoc."""
    ref_docx = BOOK_DIR / "reference.docx"
    raw_output = OUTPUT_DIR / "structural-fuzzing-book-raw.docx"

    cmd = [
        "pandoc",
        str(combined_md),
        "--from",
        "markdown+tex_math_dollars+pipe_tables+fenced_code_blocks+fenced_code_attributes+raw_tex",
        "--to",
        "docx",
        "--reference-doc",
        str(ref_docx),
        "--toc",
        "--toc-depth=3",
        "--resource-path",
        str(BOOK_DIR),
        "--output",
        str(raw_output),
    ]

    print("Running pandoc...")
    print(f"  cmd: {' '.join(cmd[:6])}...")
    result = subprocess.run(cmd, capture_output=True, text=True, cwd=str(BOOK_DIR))

    if result.returncode != 0:
        print(f"  PANDOC ERROR (exit {result.returncode}):")
        print(f"  stderr: {result.stderr[:1000]}")
        sys.exit(1)

    if result.stderr:
        # Filter out common warnings
        warnings = [
            l
            for l in result.stderr.strip().split("\n")
            if l and "Could not convert" not in l.lower()
        ]
        if warnings:
            print(f"  pandoc warnings: {len(warnings)}")
            for w in warnings[:10]:
                print(f"    {w}")

    print(f"  Raw docx: {raw_output} ({raw_output.stat().st_size / 1024:.0f} KB)")
    return raw_output


def run_postprocess(raw_docx):
    """Apply python-docx post-processing."""
    from postprocess_docx import postprocess

    final_output = OUTPUT_DIR / "structural-fuzzing-book.docx"
    postprocess(raw_docx, final_output)
    return final_output


def main():
    parser = argparse.ArgumentParser(description="Build Structural Fuzzing book")
    parser.add_argument(
        "--skip-illustrations", action="store_true", help="Skip illustration generation"
    )
    parser.add_argument(
        "--skip-pandoc", action="store_true", help="Skip pandoc conversion"
    )
    parser.add_argument(
        "--skip-postprocess",
        action="store_true",
        help="Skip python-docx post-processing",
    )
    args = parser.parse_args()

    IMAGES_DIR.mkdir(exist_ok=True)
    OUTPUT_DIR.mkdir(exist_ok=True)

    # Phase A: Generate illustrations
    if not args.skip_illustrations:
        print("=" * 60)
        print("PHASE A: Generating illustrations")
        print("=" * 60)
        from generate_illustrations import generate_all

        generate_all()
    else:
        print("Skipping illustrations (--skip-illustrations)")

    # Phase B: Pandoc conversion
    if not args.skip_pandoc:
        print("=" * 60)
        print("PHASE B: Pandoc conversion")
        print("=" * 60)
        create_reference_docx()
        combined_md = concatenate_chapters()
        run_pandoc(combined_md)
    else:
        print("Skipping pandoc (--skip-pandoc)")

    # Phase C: Post-processing
    if not args.skip_postprocess:
        print("=" * 60)
        print("PHASE C: Post-processing")
        print("=" * 60)
        raw = OUTPUT_DIR / "structural-fuzzing-book-raw.docx"
        if not raw.exists():
            print(f"ERROR: {raw} not found. Run without --skip-pandoc first.")
            sys.exit(1)
        final = run_postprocess(raw)
        size_mb = final.stat().st_size / (1024 * 1024)
        print(f"\nFinal output: {final} ({size_mb:.1f} MB)")
    else:
        print("Skipping post-processing (--skip-postprocess)")

    print("\nDone!")


if __name__ == "__main__":
    main()
