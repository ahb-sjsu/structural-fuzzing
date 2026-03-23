"""Post-process pandoc-generated docx for professional polish.

Applies: code block shading, table styling, epigraph formatting,
part title pages, and image caption styling.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from pathlib import Path


# Colors (matching generate_illustrations.py)
DARK_NAVY   = RGBColor(0x1A, 0x1A, 0x2E)
MEDIUM_BLUE = RGBColor(0x2C, 0x3E, 0x6B)
ACCENT_GOLD = RGBColor(0xD4, 0xA8, 0x43)
TEXT_DARK    = RGBColor(0x33, 0x33, 0x33)
GRAY_55     = RGBColor(0x55, 0x55, 0x55)
WHITE        = RGBColor(0xFF, 0xFF, 0xFF)


def shade_code_blocks(doc):
    """Add background shading and Consolas font to code blocks."""
    count = 0
    for para in doc.paragraphs:
        if para.style.name in ("Source Code", "Verbatim Char"):
            pPr = para.paragraph_format.element.get_or_add_pPr()
            # Background shading
            existing = pPr.findall(qn("w:shd"))
            for e in existing:
                pPr.remove(e)
            shading = parse_xml(
                f'<w:shd {nsdecls("w")} w:fill="F5F5F5" w:val="clear"/>'
            )
            pPr.append(shading)

            # Left border (thin blue line)
            pBdr = pPr.find(qn("w:pBdr"))
            if pBdr is None:
                pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}/>')
                pPr.append(pBdr)
            left_bdr = parse_xml(
                f'<w:left {nsdecls("w")} w:val="single" w:sz="12" '
                f'w:space="4" w:color="2C3E6B"/>'
            )
            existing_left = pBdr.findall(qn("w:left"))
            for e in existing_left:
                pBdr.remove(e)
            pBdr.append(left_bdr)

            # Font
            for run in para.runs:
                run.font.name = "Consolas"
                run.font.size = Pt(9)
            count += 1
    return count


def style_epigraphs(doc):
    """Style epigraphs: first italic paragraph after each Heading 1."""
    count = 0
    paras = doc.paragraphs
    for i, para in enumerate(paras):
        if para.style.name == "Heading 1":
            # Look at the next few paragraphs for an epigraph
            for j in range(i + 1, min(i + 4, len(paras))):
                cand = paras[j]
                text = cand.text.strip()
                if not text:
                    continue
                # Epigraph: starts with a quote mark and has italic runs
                has_italic = any(r.italic for r in cand.runs if r.italic is not None)
                is_quote = text.startswith(("\u201c", '"', "\u2018", "'", "\u201e"))
                if has_italic and is_quote:
                    cand.paragraph_format.left_indent = Inches(0.5)
                    cand.paragraph_format.right_indent = Inches(0.5)
                    cand.paragraph_format.space_before = Pt(6)
                    cand.paragraph_format.space_after = Pt(12)
                    for run in cand.runs:
                        run.font.size = Pt(10.5)
                        run.font.color.rgb = GRAY_55
                    count += 1
                break  # Only style the first candidate per chapter
    return count


def style_tables(doc):
    """Apply professional styling to tables."""
    count = 0
    for table in doc.tables:
        tbl = table._tbl
        # Table-wide borders
        tblPr = tbl.find(qn("w:tblPr"))
        if tblPr is None:
            tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
            tbl.insert(0, tblPr)

        borders_xml = (
            f'<w:tblBorders {nsdecls("w")}>'
            f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="AAAAAA"/>'
            f'  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="AAAAAA"/>'
            f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'  <w:insideH w:val="single" w:sz="2" w:space="0" w:color="CCCCCC"/>'
            f'  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'</w:tblBorders>'
        )
        existing = tblPr.findall(qn("w:tblBorders"))
        for e in existing:
            tblPr.remove(e)
        tblPr.append(parse_xml(borders_xml))

        # Style header row (first row)
        if table.rows:
            for cell in table.rows[0].cells:
                tc = cell._tc
                tcPr = tc.find(qn("w:tcPr"))
                if tcPr is None:
                    tcPr = parse_xml(f'<w:tcPr {nsdecls("w")}/>')
                    tc.insert(0, tcPr)
                shd = parse_xml(
                    f'<w:shd {nsdecls("w")} w:fill="1A1A2E" w:val="clear"/>'
                )
                existing_shd = tcPr.findall(qn("w:shd"))
                for e in existing_shd:
                    tcPr.remove(e)
                tcPr.append(shd)

                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.color.rgb = WHITE
                        run.font.bold = True
                        run.font.size = Pt(10)

            # Alternating row shading
            for i, row in enumerate(table.rows[1:], 1):
                if i % 2 == 0:
                    for cell in row.cells:
                        tc = cell._tc
                        tcPr = tc.find(qn("w:tcPr"))
                        if tcPr is None:
                            tcPr = parse_xml(f'<w:tcPr {nsdecls("w")}/>')
                            tc.insert(0, tcPr)
                        shd = parse_xml(
                            f'<w:shd {nsdecls("w")} w:fill="F5F5F5" w:val="clear"/>'
                        )
                        existing_shd = tcPr.findall(qn("w:shd"))
                        for e in existing_shd:
                            tcPr.remove(e)
                        tcPr.append(shd)

        count += 1
    return count


def style_image_captions(doc):
    """Style image captions: centered, italic, smaller, gray."""
    count = 0
    for i, para in enumerate(doc.paragraphs):
        if para.style.name in ("Image Caption", "Caption", "Figure"):
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.size = Pt(9)
                run.font.italic = True
                run.font.color.rgb = GRAY_55
            count += 1
        # Also detect pandoc-generated captions (paragraph right after an image)
        elif para.text.startswith("Figure ") and len(para.text) < 200:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.size = Pt(9)
                run.font.italic = True
                run.font.color.rgb = GRAY_55
            count += 1
    return count


def style_part_dividers(doc):
    """Format Part title pages (centered, large, with page break)."""
    count = 0
    for para in doc.paragraphs:
        text = para.text.strip()
        if text.startswith("Part ") and (
            "Foundations" in text or "Algorithms" in text or
            "Patterns" in text or "Systems" in text or "Appendices" in text
        ):
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_before = Pt(200)
            para.paragraph_format.space_after = Pt(24)
            for run in para.runs:
                run.font.size = Pt(28)
                run.font.color.rgb = DARK_NAVY
                run.font.bold = True
            count += 1
    return count


def postprocess(input_path, output_path):
    """Apply all post-processing steps."""
    print(f"Post-processing {input_path}...")
    doc = Document(str(input_path))

    n_code = shade_code_blocks(doc)
    print(f"  Styled {n_code} code block paragraphs")

    n_epi = style_epigraphs(doc)
    print(f"  Styled {n_epi} epigraphs")

    n_tbl = style_tables(doc)
    print(f"  Styled {n_tbl} tables")

    n_cap = style_image_captions(doc)
    print(f"  Styled {n_cap} image captions")

    n_part = style_part_dividers(doc)
    print(f"  Styled {n_part} part dividers")

    doc.save(str(output_path))
    print(f"  Saved: {output_path}")


if __name__ == "__main__":
    import sys
    if len(sys.argv) != 3:
        print("Usage: python postprocess_docx.py <input.docx> <output.docx>")
        sys.exit(1)
    postprocess(Path(sys.argv[1]), Path(sys.argv[2]))
